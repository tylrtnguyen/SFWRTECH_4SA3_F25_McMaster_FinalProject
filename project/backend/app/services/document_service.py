"""
Document Processing Service
Handles text extraction from various file formats (PDF, DOC, DOCX, TXT)
"""

import io
import logging
import tempfile
import os
from typing import Dict, Any, Optional, Tuple
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False
    textract = None

logger = logging.getLogger(__name__)

class DocumentService:
    """Service for extracting text from various document formats"""

    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc',
        'text/plain': 'txt',
    }

    MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB

    @staticmethod
    def validate_file(file_content: bytes, filename: str, mime_type: str) -> Tuple[bool, str]:
        """Validate file type and size"""
        if len(file_content) > DocumentService.MAX_FILE_SIZE:
            return False, "File size exceeds 20MB limit"

        if mime_type not in DocumentService.SUPPORTED_MIME_TYPES:
            return False, f"Unsupported file type: {mime_type}. Supported: PDF, DOC, DOCX, TXT"

        return True, "Valid"

    @staticmethod
    def extract_text(file_content: bytes, mime_type: str, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from various document formats"""
        try:
            metadata = {
                'pages': 1,
                'characters': 0,
                'words': 0,
                'extraction_method': 'unknown',
                'filename': filename,
                'mime_type': mime_type
            }

            if mime_type == 'application/pdf' and PDF_AVAILABLE:
                text, meta = DocumentService._extract_pdf_text(file_content)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              'application/msword'] and DOCX_AVAILABLE:
                text, meta = DocumentService._extract_docx_text(file_content, mime_type)
            elif mime_type == 'text/plain':
                text, meta = DocumentService._extract_plain_text(file_content)
            else:
                # Fallback to textract for other formats if available
                if TEXTRACT_AVAILABLE:
                    text, meta = DocumentService._extract_with_textract(file_content, filename)
                else:
                    raise Exception("No suitable text extraction library available for this file type")

            metadata.update(meta)
            metadata['characters'] = len(text)
            metadata['words'] = len(text.split()) if text.strip() else 0

            return text, metadata

        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {str(e)}")
            raise Exception(f"Failed to extract text from {filename}: {str(e)}")

    @staticmethod
    def _extract_pdf_text(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF files using PyPDF2"""
        if not PDF_AVAILABLE:
            raise Exception("PyPDF2 not available for PDF processing")

        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {str(e)}")
                    continue

            if not text.strip():
                raise Exception("No readable text found in PDF")

            return text.strip(), {
                'pages': len(pdf_reader.pages),
                'extraction_method': 'PyPDF2',
                'readable_pages': sum(1 for page in pdf_reader.pages if page.extract_text().strip())
            }
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise Exception(f"PDF processing failed: {str(e)}")

    @staticmethod
    def _extract_docx_text(content: bytes, mime_type: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX files using python-docx"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available for DOCX processing")

        try:
            doc = Document(io.BytesIO(content))
            text = ""

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "\n"

            if not text.strip():
                raise Exception("No readable text found in document")

            return text.strip(), {
                'extraction_method': 'python-docx',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise Exception(f"DOCX processing failed: {str(e)}")

    @staticmethod
    def _extract_plain_text(content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text = None

            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise Exception("Unable to decode text file with supported encodings")

            return text, {
                'extraction_method': 'plain_text',
                'encoding': encoding
            }
        except Exception as e:
            logger.error(f"Plain text extraction failed: {str(e)}")
            raise Exception(f"Text file processing failed: {str(e)}")

    @staticmethod
    def _extract_with_textract(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text using textract library as fallback"""
        if not TEXTRACT_AVAILABLE:
            raise Exception("textract not available for text extraction")

        temp_path = None
        try:
            # Create temporary file
            suffix = Path(filename).suffix or '.tmp'
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name

            # Extract text using textract
            text = textract.process(temp_path).decode('utf-8', errors='ignore')

            if not text.strip():
                raise Exception("No text extracted by textract")

            return text.strip(), {
                'extraction_method': 'textract',
                'temp_file_used': True
            }
        except Exception as e:
            logger.error(f"Textract extraction failed: {str(e)}")
            raise Exception(f"Advanced text extraction failed: {str(e)}")
        finally:
            # Clean up temporary file
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception as e:
                    logger.warning(f"Failed to clean up temp file {temp_path}: {str(e)}")

    @staticmethod
    def get_file_info(content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Get basic file information without extracting text"""
        return {
            'filename': filename,
            'size': len(content),
            'mime_type': mime_type,
            'extension': Path(filename).suffix.lower(),
            'supported': mime_type in DocumentService.SUPPORTED_MIME_TYPES
        }
